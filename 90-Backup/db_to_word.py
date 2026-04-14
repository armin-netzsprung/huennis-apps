import os
import django
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- DJANGO SETUP ---
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'blick_dahinter_project.settings')
django.setup()

from blog.models import BlogPost

# --- HILFSFUNKTIONEN FÜR WORD ---

def restart_list_numbering(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0')
    numId = OxmlElement('w:numId'); numId.set(qn('w:val'), '1')
    numPr.append(ilvl); numPr.append(numId)
    pPr.append(numPr)

def _handle_element(element, paragraph):
    for child in element.children:
        if isinstance(child, str):
            if child.strip() or child == ' ':
                paragraph.add_run(child)
        elif child.name == 'br':
            paragraph.add_run().add_break()
        elif child.name:
            run_text = child.get_text()
            if not run_text: continue
            run = paragraph.add_run(run_text)
            tag_name = child.name.lower()
            if tag_name in ['strong', 'b']: run.bold = True
            if tag_name in ['em', 'i']: run.italic = True

def parse_html_to_docx(html_content, doc):
    if not html_content: return
    soup = BeautifulSoup(html_content, 'html.parser')
    last_list_parent = None
    
    for tag in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'ul', 'ol', 'div', 'li']):
        # Verhindert Doppelverarbeitung
        if any(parent.name == 'div' and 'bg-sage-50' in parent.get('class', []) for parent in tag.parents):
            continue
        if any(parent.name == 'li' for parent in tag.parents) and tag.name == 'p':
            continue

        if tag.name in ['h1', 'h2', 'h3', 'h4', 'h5']:
            level = int(tag.name[1]) if int(tag.name[1]) <= 3 else 3
            p = doc.add_heading('', level=level)
            _handle_element(tag, p)
        elif tag.name == 'li':
            is_ordered = (tag.parent.name == 'ol')
            p = doc.add_paragraph(style='List Number' if is_ordered else 'List Bullet')
            if is_ordered and last_list_parent != tag.parent:
                restart_list_numbering(p)
                last_list_parent = tag.parent
            _handle_element(tag, p)
        elif tag.name == 'div' and 'bg-sage-50' in tag.get('class', []):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _handle_element(tag, p)
        elif tag.name == 'p':
            if tag.get_text(strip=True):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _handle_element(tag, p)

# --- HAUPTFUNKTION ---

def run_export(category_name):
    # Erstmal prüfen, was überhaupt da ist
    total_in_db = BlogPost.objects.count()
    print(f"Datenbank-Check: {total_in_db} Beiträge insgesamt gefunden.")

    # Filtern nach Kategorie (exakter Name oder Teilwort)
    posts = BlogPost.objects.filter(category__name__icontains=category_name).order_by('part_title', 'chapter_title')
    
    post_count = posts.count()
    if post_count == 0:
        print(f"Fehler: Keine Beiträge für '{category_name}' gefunden.")
        return

    print(f"Starte Export von {post_count} Kapiteln...")
    doc = Document()
    doc.add_heading(f'Manuskript: {category_name}', 0)

    current_part = None
    for i, post in enumerate(posts, 1):
        print(f"[{i}/{post_count}] Verarbeite: {post.chapter_title}")
        
        # Neuer Teil = Neue Seite
        if post.part_title != current_part:
            if current_part is not None:
                doc.add_page_break()
            doc.add_heading(post.part_title or "Inhalt", level=1)
            current_part = post.part_title

        # Kapitel & Inhalt
        doc.add_heading(post.chapter_title, level=2)
        if post.chapter_subtitle:
            doc.add_paragraph().add_run(post.chapter_subtitle).italic = True
            
        parse_html_to_docx(post.introduction, doc)
        parse_html_to_docx(post.content, doc)

    output_file = f"Manuskript_{category_name.replace(' ', '_')}.docx"
    doc.save(output_file)
    print(f"\nERFOLG! Datei gespeichert als: {output_file}")

if __name__ == "__main__":
    # In der Praxis auf dem Server aufrufen:
    run_export("Sucht (Alkohol)")
    