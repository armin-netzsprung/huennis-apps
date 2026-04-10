import sys
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def restart_list_numbering(paragraph):
    """Fügt dem Paragraphen XML-Eigenschaften hinzu, um die Nummerierung bei 1 zu starten."""
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '1') # Setzt eine Basis-ID; Word erstellt bei neuen Blöcken meist neue
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)

def _handle_element(element, paragraph):
    """Verarbeitet Text, Zeilenumbrüche und Inline-Formatierung."""
    for child in element.children:
        if isinstance(child, str):
            if child.strip() or child == ' ':
                paragraph.add_run(child)
        elif child.name == 'br':
            paragraph.add_run().add_break()
        elif child.name:
            run_text = child.get_text()
            if not run_text: continue
            run = paragraph.add_run(run_text)
            tag_name = child.name.lower()
            style_str = str(child.get('style', '')).lower()
            if tag_name in ['strong', 'b'] or 'bold' in style_str:
                run.bold = True
            if tag_name in ['em', 'i'] or 'italic' in style_str:
                run.italic = True

def parse_html_to_docx(html_content, doc):
    soup = BeautifulSoup(html_content, 'html.parser')
    last_list_parent = None
    
    for tag in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'ul', 'ol', 'div', 'li']):
        # Parent-Check gegen Doppelverarbeitung
        if any(parent.name == 'div' and 'bg-sage-50' in parent.get('class', []) for parent in tag.parents):
            continue
        if any(parent.name == 'li' for parent in tag.parents) and tag.name == 'p':
            continue

        # 1. Überschriften
        if tag.name in ['h1', 'h2', 'h3', 'h4', 'h5']:
            level = int(tag.name[1]) if int(tag.name[1]) <= 3 else 3
            p = doc.add_heading('', level=level)
            _handle_element(tag, p)

        # 2. Listenpunkte
        elif tag.name == 'li':
            # Finde heraus, ob wir in einer <ol> oder <ul> sind
            list_type = tag.parent.name 
            is_ordered = (list_type == 'ol')
            
            p = doc.add_paragraph(style='List Number' if is_ordered else 'List Bullet')
            
            # NEU: Prüfen, ob eine neue geordnete Liste beginnt
            if is_ordered:
                if last_list_parent != tag.parent:
                    # Dies ist der erste Punkt einer neuen Liste -> Neustart erzwingen
                    restart_list_numbering(p)
                    last_list_parent = tag.parent
            
            _handle_element(tag, p)

        # 3. Zitat-Boxen
        elif tag.name == 'div' and 'bg-sage-50' in tag.get('class', []):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _handle_element(tag, p)

        # 4. Normale Absätze
        elif tag.name == 'p':
            text_content = tag.get_text(strip=True)
            if text_content:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _handle_element(tag, p)

def main():
    if len(sys.argv) < 2:
        print("Benutzung: python test_converter.py input.html")
        return
    input_file = sys.argv[1]
    with open(input_file, 'r', encoding='utf-8') as f:
        html_data = f.read()
    doc = Document()
    parse_html_to_docx(html_data, doc)
    doc.save("test_output_fixed_v2.docx")
    print("Erfolg! Nummerierung wurde zurückgesetzt.")

if __name__ == "__main__":
    main()